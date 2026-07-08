#!/usr/bin/env python3
"""
Convert DOCX/DOC files to Markdown for raw/ ingestion.
Extracts text, tables, AND embedded images (saved to ./assets/ for Obsidian browsing;
note: AI cannot read standalone image files — image content must be described in text).

Note: PDF is NOT converted — Read tool renders PDF pages visually (charts,
formulas, screenshots all visible within the page context), which is richer
than text extraction.

Usage:
    python scripts/convert-to-markdown.py input.docx  → output: input.md + assets/
    python scripts/convert-to-markdown.py input.doc   → output: input.md + assets/

Dependencies: python-docx, LibreOffice soffice (for .doc only)
Install: pip install python-docx
"""

import argparse
import base64
import io
import os
import re
import subprocess
import sys
from pathlib import Path


# ---------------------------------------------------------------------------
# Image extraction helpers
# ---------------------------------------------------------------------------

def _extract_docx_images(docx_path: str, assets_dir: Path) -> dict[str, Path]:
    """Extract all images from a docx file. Returns {rId: relative_path}."""
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    doc = Document(docx_path)
    image_map: dict[str, Path] = {}
    base_name = Path(docx_path).stem

    for idx, rel in enumerate(doc.part.rels.values()):
        if "image" not in rel.reltype:
            continue
        image_data = rel.target_part.blob

        # Determine extension from content type
        ext_map = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/tiff": ".tiff",
            "image/x-emf": ".emf",
            "image/x-wmf": ".wmf",
            "image/svg+xml": ".svg",
        }
        content_type = rel.target_part.content_type
        ext = ext_map.get(content_type, ".png")

        img_name = f"{base_name}-img{idx + 1}{ext}"
        img_path = assets_dir / img_name
        img_path.write_bytes(image_data)
        image_map[rel.rId] = Path("assets") / img_name

    return image_map


def _find_images_in_paragraphs(docx_path: str, image_map: dict[str, Path]) -> list[str]:
    """Walk paragraphs and runs, inserting image references at the right spots."""
    from docx import Document
    from docx.oxml.ns import qn

    IMAGE_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    doc = Document(docx_path)
    lines = []
    seen_images: set[str] = set()

    for para in doc.paragraphs:
        para_has_content = False
        para_parts: list[str] = []
        seen_drawings = False

        # Check paragraph style for heading
        style_name = para.style.name if para.style else ""
        heading_level = _heading_level(style_name)

        for run in para.runs:
            # Check for inline images in this run
            drawings = run._element.findall(f"{IMAGE_NS}drawing")
            for drawing in drawings:
                blips = drawing.findall(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                )
                for blip in blips:
                    embed = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if embed and embed in image_map:
                        img_rel_path = image_map[embed]
                        lines.append(
                            f"\n![image]({img_rel_path.as_posix()})\n"
                        )
                        seen_images.add(embed)
                        para_has_content = True

            text = run.text
            if text:
                para_parts.append(text)
                para_has_content = True

        if para_parts:
            text = "".join(para_parts)
            if heading_level:
                lines.append(f"{'#' * heading_level} {text}")
            else:
                lines.append(text)

        if not para_has_content:
            lines.append("")
        else:
            lines.append("")  # paragraph separator

        # Fallback: unreferenced images after first paragraph
        if not seen_drawings and image_map and not seen_images:
            for rId, img_path in image_map.items():
                if rId not in seen_images:
                    lines.append(
                        f"\n![image]({img_path.as_posix()})\n"
                    )
                    seen_images.add(rId)

    return lines


# ---------------------------------------------------------------------------
# DOCX conversion
# ---------------------------------------------------------------------------

def _heading_level(style_name: str) -> int:
    """Extract heading level from Word style name."""
    for pattern, level in [
        (r"^Heading\s*1\b", 1),
        (r"^Heading\s*2\b", 2),
        (r"^Heading\s*3\b", 3),
        (r"^Heading\s*4\b", 4),
        (r"^Heading\s*5\b", 5),
        (r"^Heading\s*6\b", 6),
        (r"^Title\b", 1),
        (r"^Subtitle\b", 2),
    ]:
        if re.search(pattern, style_name, re.IGNORECASE):
            return level
    return 0


def docx_to_markdown(docx_path: str, output_path: str) -> int:
    """Convert .docx to Markdown with embedded images extracted. Returns image count."""
    from docx import Document

    output_dir = Path(output_path).parent
    assets_dir = output_dir / "assets"
    assets_dir.mkdir(parents=True, exist_ok=True)

    # Step 1: Extract images
    image_map = _extract_docx_images(docx_path, assets_dir)

    # Step 2: Build text with image references inline
    lines = _find_images_in_paragraphs(docx_path, image_map)

    # Step 3: Tables (after paragraphs)
    doc = Document(docx_path)
    if doc.tables:
        lines.append("")
        for table in doc.tables:
            lines.append("")
            for row_idx, row in enumerate(table.rows):
                cells = [
                    cell.text.strip().replace("\n", " ")
                    for cell in row.cells
                ]
                lines.append("| " + " | ".join(cells) + " |")
                if row_idx == 0:
                    lines.append(
                        "|" + "|".join([" --- " for _ in cells]) + "|"
                    )
            lines.append("")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return len(image_map)


# ---------------------------------------------------------------------------
# DOC conversion (via LibreOffice → DOCX)
# ---------------------------------------------------------------------------

def _find_soffice() -> str:
    """Locate LibreOffice soffice executable."""
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/usr/local/bin/soffice",
        "soffice",
    ]
    for candidate in candidates:
        try:
            subprocess.run(
                [candidate, "--version"],
                capture_output=True,
                check=True,
                timeout=10,
            )
            return candidate
        except (
            FileNotFoundError,
            subprocess.TimeoutExpired,
            subprocess.CalledProcessError,
        ):
            continue

    raise FileNotFoundError(
        "LibreOffice not found. Install it first, then re-run this script."
    )


def doc_to_markdown(doc_path: str, output_path: str) -> int:
    """Convert .doc → .docx via LibreOffice, then to Markdown with images."""
    soffice = _find_soffice()
    output_dir = Path(output_path).parent
    tmp_dir = output_dir / ".tmp_convert"
    tmp_dir.mkdir(exist_ok=True)

    print("Converting .doc → .docx via LibreOffice...")
    result = subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(tmp_dir),
            doc_path,
        ],
        capture_output=True,
        text=True,
        timeout=120,
    )

    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

    docx_files = list(tmp_dir.glob("*.docx"))
    if not docx_files:
        raise FileNotFoundError(f"No .docx produced in {tmp_dir}")

    docx_path = docx_files[0]
    print(f"Generated: {docx_path.name}")

    img_count = docx_to_markdown(str(docx_path), output_path)

    # Cleanup temp files
    docx_path.unlink()
    try:
        tmp_dir.rmdir()
    except OSError:
        pass

    return img_count


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert DOCX/DOC files to Markdown (with images)"
    )
    parser.add_argument(
        "input", help="Input file (.docx or .doc)"
    )
    parser.add_argument(
        "-o",
        "--output",
        help="Output path (default: same name with .md)",
    )
    args = parser.parse_args()

    input_path = args.input
    if not os.path.exists(input_path):
        print(f"Error: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    output_path = args.output or str(Path(input_path).with_suffix(".md"))
    ext = Path(input_path).suffix.lower()

    print(f"Converting: {input_path} → {output_path}")

    if ext == ".docx":
        img_count = docx_to_markdown(input_path, output_path)
        print(
            f"Done. {img_count} image(s) extracted → assets/"
            if img_count
            else "Done."
        )
    elif ext == ".doc":
        img_count = doc_to_markdown(input_path, output_path)
        print(
            f"Done. {img_count} image(s) extracted → assets/"
            if img_count
            else "Done."
        )
    elif ext == ".pdf":
        print(
            "PDF files are read natively by the Read tool (visual rendering). "
            "No conversion needed — move it directly into raw/.",
            file=sys.stderr,
        )
        sys.exit(0)
    else:
        print(
            f"Error: Unsupported format '{ext}'. Supported: .docx, .doc",
            file=sys.stderr,
        )
        sys.exit(1)


if __name__ == "__main__":
    main()
